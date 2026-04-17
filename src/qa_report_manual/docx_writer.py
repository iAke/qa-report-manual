"""Render SiteSnapshot to styled .docx (QA report + user manual)."""
from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm, Inches

from .models import PageSnapshot, SiteSnapshot


# ---------- styling helpers ----------

BRAND = RGBColor(0x0B, 0x5F, 0xA5)   # deep blue
ACCENT = RGBColor(0xE8, 0xF1, 0xFA)  # light blue row shading
TEXT_MUTED = RGBColor(0x55, 0x55, 0x55)
HEADER_FILL = "0B5FA5"
ZEBRA_FILL = "F2F6FB"


def _set_cell_shading(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _set_cell_text(cell, text: str, *, bold: bool = False, color: RGBColor | None = None,
                   size: int = 10, align=None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _configure_base_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for name, size, color, bold in [
        ("Heading 1", 22, BRAND, True),
        ("Heading 2", 16, BRAND, True),
        ("Heading 3", 13, BRAND, True),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = bold
        st.font.color.rgb = color


def _add_table(doc: Document, headers: list[str], rows: list[list[str]],
               col_widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    # header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell_text(hdr[i], h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=10)
        _set_cell_shading(hdr[i], HEADER_FILL)

    # body rows
    for r, row in enumerate(rows, start=1):
        cells = table.rows[r].cells
        zebra = (r % 2 == 0)
        for i, val in enumerate(row):
            _set_cell_text(cells[i], val if val is not None else "", size=10)
            if zebra:
                _set_cell_shading(cells[i], ZEBRA_FILL)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)


def _add_kv_table(doc: Document, pairs: list[tuple[str, str]]) -> None:
    t = doc.add_table(rows=len(pairs), cols=2)
    t.style = "Light Grid Accent 1"
    t.autofit = False
    for i, (k, v) in enumerate(pairs):
        _set_cell_text(t.rows[i].cells[0], k, bold=True, size=10)
        _set_cell_shading(t.rows[i].cells[0], ZEBRA_FILL)
        _set_cell_text(t.rows[i].cells[1], v, size=10)
        t.rows[i].cells[0].width = Cm(5.5)
        t.rows[i].cells[1].width = Cm(11)


def _add_figure(doc: Document, image_path: str, caption: str, width_in: float = 6.0,
                number: int | None = None) -> None:
    """Embed an image scaled to width_in inches with a numbered italic caption."""
    import os
    if not image_path or not os.path.exists(image_path):
        return
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=Inches(width_in))
    except Exception:
        return
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label = f"Figure {number}. {caption}" if number is not None else caption
    r = cap.add_run(label)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = TEXT_MUTED


def _muted(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.color.rgb = TEXT_MUTED
    r.font.size = Pt(10)


def _bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def _number(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def _add_toc_field(doc: Document) -> None:
    """Insert a real Word TOC field (user must right-click → Update Field to populate)."""
    p = doc.add_paragraph()
    run = p.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click → Update Field to populate the Table of Contents."
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    run._r.append(placeholder)
    run._r.append(fld_char3)


def _verdict_for(site) -> tuple[str, str]:
    """Return (verdict, color hex) based on auto-findings severity mix."""
    sev = [f.severity for f in site.findings]
    if any(s == "S1" for s in sev):
        return "FAIL — critical issues block release", "B3261E"
    if sum(1 for s in sev if s == "S2") >= 2:
        return "CONDITIONAL — major issues need remediation", "A56400"
    if sev:
        return "PASS WITH NOTES — minor/moderate issues", "16794A"
    return "PASS — no automated red flags (manual QA still required)", "16794A"


def _figure_counter():
    n = {"i": 0}
    def nxt():
        n["i"] += 1
        return n["i"]
    return nxt


def _title_block(doc: Document, title: str, subtitle_pairs: list[tuple[str, str]]) -> None:
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = t.add_run(title)
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = BRAND

    # colored underline bar
    bar = doc.add_paragraph()
    bar_run = bar.add_run("▬" * 30)
    bar_run.font.color.rgb = BRAND
    bar_run.font.size = Pt(10)

    _add_kv_table(doc, subtitle_pairs)
    doc.add_paragraph()


# ---------- QA report ----------

def write_qa_report_docx(site: SiteSnapshot, out_path: Path) -> None:
    doc = Document()
    _configure_base_styles(doc)

    now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    pages = site.pages
    failed = [p for p in pages if p.load_error]
    ok = [p for p in pages if not p.load_error]

    total_imgs = sum(p.total_images for p in ok)
    total_missing_alt = sum(p.images_without_alt for p in ok)
    total_forms = sum(len(p.forms) for p in ok)
    total_console_errs = sum(len(p.console_errors) for p in ok)
    multi_h1 = [p for p in ok if sum(1 for lvl, _ in p.headings if lvl == 1) > 1]
    no_h1 = [p for p in ok if not any(lvl == 1 for lvl, _ in p.headings)]
    no_meta = [p for p in ok if not p.meta_description]
    no_lang = [p for p in ok if not p.lang]

    _title_block(
        doc,
        "QA Report — Site Crawl",
        [
            ("Entry URL", site.entry_url),
            ("Pages crawled", str(len(pages))),
            ("Generated", now),
            ("Tool", "qa-report-manual (Playwright + python-docx)"),
            ("Standard", "Structure follows ISO/IEC/IEEE 29119-3 §7 Test Report"),
        ],
    )

    # Executive summary + verdict
    verdict, color = _verdict_for(site)
    doc.add_heading("Executive summary", level=1)
    vp = doc.add_paragraph()
    vr = vp.add_run(verdict)
    vr.bold = True
    vr.font.size = Pt(13)
    vr.font.color.rgb = RGBColor.from_string(color)
    doc.add_paragraph(
        f"Automated structural pass across {len(pages)} same-site page(s) reachable from the entry URL. "
        f"{len(failed)} page(s) failed to load. "
        f"{len(site.findings)} auto-derived finding(s) — see §4 for details."
    )

    # Severity snapshot
    sev_counts = {"S1": 0, "S2": 0, "S3": 0, "S4": 0}
    for f in site.findings:
        sev_counts[f.severity] = sev_counts.get(f.severity, 0) + 1
    _add_table(
        doc,
        ["Severity", "Count", "Meaning"],
        [
            ["S1", str(sev_counts["S1"]), "Production down / data loss / blocker"],
            ["S2", str(sev_counts["S2"]), "Major feature broken"],
            ["S3", str(sev_counts["S3"]), "Moderate, workaround exists"],
            ["S4", str(sev_counts["S4"]), "Minor / cosmetic"],
        ],
        col_widths=[3, 2, 11],
    )

    # Table of contents (Word field)
    doc.add_heading("Table of contents", level=1)
    _add_toc_field(doc)

    # 1. Purpose
    doc.add_heading("1. Purpose and scope", level=1)
    doc.add_paragraph(
        "Automated structural pass across same-site pages reachable from the entry URL. "
        "This is not a full test run — it produces inputs for manual QA and a defect log skeleton."
    )

    # 2. Environment
    doc.add_heading("2. Test environment", level=1)
    _add_kv_table(doc, [
        ("Browser / engine", "Chromium (Playwright, headless)"),
        ("Viewport", "1280×720"),
        ("Network", "_fill in_"),
        ("Test data / roles", "_fill in_"),
    ])

    # 3. Overall summary
    doc.add_heading("3. Overall summary", level=1)
    _add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Pages crawled", str(len(pages))],
            ["Pages failed to load", str(len(failed))],
            ["Total forms detected", str(total_forms)],
            ["Total images", str(total_imgs)],
            ["Images missing alt", str(total_missing_alt)],
            ["Console errors (sum)", str(total_console_errs)],
            ["Pages with multiple H1", str(len(multi_h1))],
            ["Pages with no H1", str(len(no_h1))],
            ["Pages missing meta description", str(len(no_meta))],
            ["Pages missing <html lang>", str(len(no_lang))],
        ],
        col_widths=[8, 4],
    )

    # Findings (severity-rated)
    doc.add_heading("4. Auto-derived findings", level=1)
    if site.findings:
        rows = [[f.severity, f.area, f.message, f.reference or "-", str(len(f.pages))]
                for f in sorted(site.findings, key=lambda x: x.severity)]
        _add_table(
            doc,
            ["Severity", "Area", "Finding", "Reference", "Pages"],
            rows,
            col_widths=[1.6, 2.5, 7, 3, 1.2],
        )
    else:
        _muted(doc, "No automated red flags detected; manual QA still required.")

    # Performance snapshot
    doc.add_heading("Performance snapshot", level=2)
    perf_rows = []
    for p in ok:
        perf_rows.append([
            p.final_url or p.url,
            f"{p.perf.load_ms} ms",
            f"{p.perf.dom_content_loaded_ms} ms",
            f"{p.perf.transfer_bytes / 1024:.0f} KB",
            str(p.perf.resource_count),
        ])
    if perf_rows:
        _add_table(
            doc,
            ["URL", "Load", "DCL", "Transfer", "Resources"],
            perf_rows,
            col_widths=[7, 2, 2, 2.5, 2],
        )

    if site.sitemap_count:
        _muted(doc, f"Sitemap.xml detected — {site.sitemap_count} URL(s) seeded for discovery.")

    # 5. Per-page
    doc.add_heading("5. Per-page overview", level=1)
    rows = []
    for p in pages:
        if p.load_error:
            rows.append([p.final_url or p.url, "ERROR", "-", "-", "-", "-", p.load_error[:80]])
        else:
            h1c = sum(1 for lvl, _ in p.headings if lvl == 1)
            rows.append([
                p.final_url or p.url, "OK", p.title or "(empty)",
                str(h1c), f"{p.images_without_alt}/{p.total_images}",
                str(len(p.forms)), str(len(p.console_errors)),
            ])
    _add_table(
        doc,
        ["URL", "Status", "Title", "H1", "Missing alt / total", "Forms", "Console errs"],
        rows,
        col_widths=[5.5, 1.5, 4.2, 1, 2.5, 1.3, 1.5],
    )

    # 6. Console errors
    doc.add_heading("6. Console errors (sample)", level=1)
    cerr_rows: list[list[str]] = []
    for p in ok:
        for e in p.console_errors[:5]:
            cerr_rows.append([p.final_url or p.url, e[:300]])
    if cerr_rows:
        _add_table(doc, ["Page", "Message"], cerr_rows, col_widths=[6, 11])
    else:
        _muted(doc, "No console errors captured during the crawl.")

    # 7. Forms
    doc.add_heading("7. Forms detected across site", level=1)
    any_form = False
    idx = 0
    for p in ok:
        for f in p.forms:
            idx += 1
            any_form = True
            doc.add_heading(f"Form {idx} — {p.final_url or p.url}", level=2)
            _muted(doc, f"{f.method} {f.action or '/'}")
            if f.fields:
                _add_table(
                    doc,
                    ["Label / name", "Type", "Required"],
                    [[x.label, x.type, "Yes" if x.required else "No"] for x in f.fields],
                    col_widths=[8, 4, 3],
                )
            else:
                _muted(doc, "No fields detected.")
    if not any_form:
        _muted(doc, "No <form> elements detected on any crawled page.")

    # 8. Checklist
    doc.add_heading("8. Manual test checklist", level=1)
    _add_table(
        doc,
        ["Area", "Test ID", "Result", "Notes"],
        [
            ["Smoke: every crawled page loads", "", "Pass / Fail / Blocked", ""],
            ["Navigation consistency across pages", "", "", ""],
            ["Forms: validation & required fields", "", "", ""],
            ["Links: no broken internal links", "", "", ""],
            ["Accessibility: landmarks, heading order, alt", "", "", ""],
            ["Responsive: mobile / tablet / desktop", "", "", ""],
            ["Cross-browser: Chrome / Firefox / Safari", "", "", ""],
            ["Security: HTTPS, cookie flags", "", "", ""],
            ["Performance: LCP / CLS spot check", "", "", ""],
        ],
        col_widths=[8, 2, 3, 4],
    )
    _muted(doc, "Severity guide: S1 production down / data loss · S2 major feature broken · S3 moderate with workaround · S4 minor / cosmetic.")

    # 9. Defect log
    doc.add_heading("9. Defect log (template)", level=1)
    _add_table(
        doc,
        ["ID", "Severity", "Page URL", "Summary", "Steps", "Expected", "Actual", "Evidence"],
        [["DEF-001", "S1–S4", "", "", "", "", "", "screenshot / HAR"]],
    )

    # 10. Sign-off
    doc.add_heading("10. Sign-off", level=1)
    _add_table(
        doc,
        ["Role", "Name", "Date", "Comments"],
        [["QA", "", "", ""], ["Product", "", "", ""], ["Engineering", "", "", ""]],
    )

    doc.save(str(out_path))


# ---------- User manual ----------

def _page_sections(p: PageSnapshot) -> list[tuple[str, list[str]]]:
    sections: list[tuple[str, list[str]]] = []
    current: str | None = None
    buf: list[str] = []
    for level, text in p.headings:
        if level == 2:
            if current is not None:
                sections.append((current, buf))
            current = text
            buf = []
        elif level == 3 and current is not None:
            buf.append(text)
    if current is not None:
        sections.append((current, buf))
    if not sections:
        h1s = [t for lvl, t in p.headings if lvl == 1]
        if h1s:
            sections = [(t, []) for t in h1s]
        else:
            sections = [(p.title or "Using this page", [])]
    return sections


def write_user_manual_docx(site: SiteSnapshot, out_path: Path) -> None:
    doc = Document()
    _configure_base_styles(doc)

    now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    pages = [p for p in site.pages if not p.load_error]
    home = pages[0] if pages else None
    site_title = (home.title if home else "") or site.entry_url
    intro = (home.meta_description if home else "") or f"This guide explains how to use {site_title}."

    _title_block(
        doc,
        f"User Manual — {site_title}",
        [
            ("Entry URL", site.entry_url),
            ("Pages covered", str(len(pages))),
            ("Generated", now),
        ],
    )

    doc.add_heading("About this manual", level=1)
    doc.add_paragraph(intro)
    _muted(doc, "This document was machine-generated from page structure. Replace placeholders with accurate product language, screenshots, and policies before publication.")

    doc.add_heading("Audience", level=1)
    _bullet(doc, "Primary: _e.g. new customers, internal staff_")
    _bullet(doc, "Prerequisites: _accounts, hardware, supported browsers_")

    doc.add_heading("Table of contents", level=1)
    _add_toc_field(doc)
    _muted(doc, "Structure follows DITA information typing: each page has Overview (concept), Steps (task), and Reference sections.")

    doc.add_heading("Quick start", level=1)
    _number(doc, f"Visit {site.entry_url}.")
    _number(doc, "Use the main navigation to reach the section you need.")
    _number(doc, "In each page section below, follow the numbered Steps.")
    _number(doc, "Verify the Expected outcome matches what you see.")

    fig = _figure_counter()
    for p in pages:
        label = p.title or p.final_url or p.url
        doc.add_heading(label, level=1)
        _muted(doc, p.final_url or p.url)

        # --- Overview (concept)
        doc.add_heading("Overview", level=2)
        if p.meta_description:
            doc.add_paragraph(p.meta_description)
        else:
            doc.add_paragraph(f"This page is part of {site.entry_url}. Use it to accomplish the tasks listed below.")

        if p.screenshot_viewport:
            _add_figure(doc, p.screenshot_viewport,
                        f"Above-the-fold view of {label}.", number=fig())
        if p.screenshot_mobile:
            _add_figure(doc, p.screenshot_mobile,
                        f"Mobile view (iPhone 14) of {label}.", width_in=3.0, number=fig())

        # --- Steps (task)
        doc.add_heading("Steps", level=2)
        sections = _page_sections(p)
        for title, bullets in sections:
            doc.add_heading(f"To {title.lower()}", level=3)
            _number(doc, f"Open {p.final_url or p.url}.")
            _number(doc, f"Locate the {title} area on the page.")
            if bullets:
                _number(doc, f"Review sub-topics: {', '.join(bullets[:6])}.")
            _number(doc, "Complete the primary action (submit, read, or select as required).")
            vp = doc.add_paragraph()
            vr = vp.add_run("Expected outcome: ")
            vr.bold = True
            vp.add_run("the action succeeds with visible confirmation (success message, navigation, or data change).")

        for idx, sp in enumerate(p.form_screenshots, start=1):
            _add_figure(doc, sp, f"Form {idx} on {label}.", width_in=5.5, number=fig())

        # --- Reference
        doc.add_heading("Reference", level=2)
        if p.buttons:
            doc.add_paragraph("Controls detected on this page:")
            _add_table(doc, ["Control label"], [[b] for b in p.buttons[:15]], col_widths=[14])
        ref_rows = [
            ["Language", p.lang or "(not set)"],
            ["Heading count (H1/H2/H3)",
             f"{sum(1 for l,_ in p.headings if l==1)} / {sum(1 for l,_ in p.headings if l==2)} / {sum(1 for l,_ in p.headings if l==3)}"],
            ["Forms", str(len(p.forms))],
            ["Images", str(p.total_images)],
            ["Load time",
             f"{p.perf.load_ms} ms ({p.perf.transfer_bytes/1024:.0f} KB)" if p.perf.load_ms else "n/a"],
        ]
        _add_table(doc, ["Attribute", "Value"], ref_rows, col_widths=[5, 12])

        if p.screenshot_fullpage:
            _add_figure(doc, p.screenshot_fullpage,
                        f"Full page of {label}.", width_in=6.0, number=fig())

    doc.add_heading("Troubleshooting", level=1)
    _add_table(
        doc,
        ["Symptom", "What to try"],
        [
            ["Page does not load", "Check network; try another browser; clear cache"],
            ["Cannot sign in", "Verify credentials; reset password; check caps lock"],
            ["Missing data", "Refresh; verify filters; contact support with URL and time"],
        ],
        col_widths=[6, 11],
    )

    doc.add_heading("Glossary", level=1)
    _add_table(doc, ["Term", "Definition"], [["_add terms_", ""]], col_widths=[5, 12])

    doc.add_heading("Support", level=1)
    _add_kv_table(doc, [
        ("Contact", "_support email / portal_"),
        ("Privacy & terms", "_link to legal pages_"),
    ])

    doc.add_heading("Feedback on this manual", level=1)
    doc.add_paragraph(
        "Help us improve this documentation. If instructions were unclear or a screenshot is stale, "
        "reply with the section title and what you expected. Feedback is reviewed for the next release."
    )
    _add_kv_table(doc, [
        ("Document version", "0.2 (draft)"),
        ("Last updated", now),
        ("Source", "Auto-generated by qa-report-manual; human review required before publication."),
    ])

    doc.save(str(out_path))
